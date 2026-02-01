import os
import logging
import requests
import json
import re
import asyncio

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload, MediaIoBaseUpload
from googleapiclient.errors import HttpError

from openai import AsyncOpenAI

from config import (
    ELEVENLABS_STT_URL,
    ELEVENLABS_API_KEY,
    ELEVENLABS_MODEL_ID,
    ELEVENLABS_TAG_AUDIO_EVENTS,
    ELEVENLABS_DIARIZE,
    GOOGLE_APPLICATION_CREDENTIALS,
    GOOGLE_CREDENTIALS_JSON,
    OPENAI_API_KEY
)

SUPPORTED_EXTENSIONS = [".mp3", ".wav", ".mp4", ".m4a", ".flac", ".ogg", ".aac"]

def replace_quotes_with_ukrainian(text: str) -> str:
    """
    Replace all double quote variations with Ukrainian quotation marks « and ».
    Alternates between opening « and closing » quotes.
    Leaves single quotes (apostrophes) untouched.
    """
    # Pattern to match all variations of double quotes
    # This includes: " (straight), " (left double), " (right double), and similar
    quote_pattern = r'[""\u201C\u201D\u201E\u201F\u2033\u2036]'

    opening = True
    def replace_quote(match):
        nonlocal opening
        if opening:
            opening = False
            return '«'
        else:
            opening = True
            return '»'

    return re.sub(quote_pattern, replace_quote, text)

def format_timestamp(seconds: float) -> str:
    """
    Convert float-second timestamp into "HH:MM:SS,mmm".
    """
    total_ms = int(seconds * 1_000)
    h = total_ms // 3_600_000
    m = (total_ms % 3_600_000) // 60_000
    s = (total_ms % 60_000) // 1_000
    ms = total_ms % 1_000
    return f"{h:02}:{m:02}:{s:02},{ms:03}"

def create_transcript(transcription_result):
    words = transcription_result.get("words", [])
    if not words:
        return transcription_result.get("text", "")

    segments = []
    current_segment = []
    current_speaker = None

    for word in words:
        speaker = word.get("speaker_id", "unknown")
        if current_speaker is None:
            current_speaker = speaker
        elif speaker != current_speaker:
            segments.append((current_speaker, current_segment))
            current_segment = []
            current_speaker = speaker
        current_segment.append(word)

    if current_segment:
        segments.append((current_speaker, current_segment))

    transcript_lines = []
    for speaker, seg_words in segments:
        start_time = seg_words[0]["start"]
        end_time = seg_words[-1]["end"]
        start_formatted = format_timestamp(start_time)
        end_formatted = format_timestamp(end_time)
        segment_text = ''.join(word["text"] for word in seg_words)
        segment_text = replace_quotes_with_ukrainian(segment_text)
        transcript_lines.append(f"{start_formatted} --> {end_formatted} - [{speaker}]")
        transcript_lines.append("")
        transcript_lines.append(segment_text)
        transcript_lines.append("")

    return "\n".join(transcript_lines)

def create_srt_from_json(transcript_json: dict,
                         max_chars: int = 40,
                         max_duration: float = 4.0) -> str:
    """
    Build an SRT file from word-level JSON:
     - Gather words until you hit:
        • sentence-ending punctuation (.,!?),
        • OR accumulated text length > max_chars,
        • OR time span > max_duration seconds
     - Then flush that buffer as one subtitle block.
    Returns the full SRT as a single string.
    """
    segments = []
    buf = []

    for w in transcript_json.get("words", []):
        if w["type"] != "word":
            continue
        buf.append(w)
        text = " ".join([x["text"].strip() for x in buf]).strip()
        duration = buf[-1]["end"] - buf[0]["start"]

        # should we cut here?
        ends_sentence = buf[-1]["text"].rstrip().endswith((".", "!", "?"))
        if ends_sentence or len(text) >= max_chars or duration >= max_duration:
            segments.append(buf)
            buf = []

    # leftover
    if buf:
        segments.append(buf)

    # now render SRT
    out_lines = []
    for idx, seg in enumerate(segments, start=1):
        start_ts = format_timestamp(seg[0]["start"])
        end_ts   = format_timestamp(seg[-1]["end"])
        line_txt = " ".join([w["text"].strip() for w in seg]).strip()

        out_lines.append(str(idx))
        out_lines.append(f"{start_ts} --> {end_ts}")
        out_lines.append(line_txt)
        out_lines.append("")  # blank line

    # Apply quote replacement to the entire SRT content after assembly
    # This ensures quotes spanning multiple segments are handled correctly
    full_srt = "\n".join(out_lines)
    return replace_quotes_with_ukrainian(full_srt)

def cleanup_temp_file(file_path):
    """Remove temporary file."""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logging.info(f"Cleaned up temporary file: {file_path}")
    except Exception as e:
        logging.error(f"Error cleaning up temporary file {file_path}: {e}")

def is_audio_or_video(fileinfo):
    """Check if the file is audio or video
    based on Slack's 'mimetype' or 'filetype'."""
    mimetype = fileinfo.get("mimetype", "")
    if mimetype.startswith("audio") or mimetype.startswith("video"):
        return True
    # Fallback: check file extension if mimetype is generic
    filename = fileinfo.get("name", "").lower()
    return filename.endswith(tuple(SUPPORTED_EXTENSIONS))

def file_too_large(fileinfo):
    size_bytes = fileinfo.get("size", 0)
    return size_bytes > 1000 * 1_000_000

def transcribe_file(file_path: str, as_srt: bool = False):
    """
    Send `file_path` to ElevenLabs.
    - Returns JSON (dict) when as_srt=False (default).
    - Returns raw SRT text (str) when as_srt=True.
    """
    headers = {"xi-api-key": ELEVENLABS_API_KEY}
    data = {
        "model_id": ELEVENLABS_MODEL_ID,
        "tag_audio_events": ELEVENLABS_TAG_AUDIO_EVENTS,
        "diarize": ELEVENLABS_DIARIZE,
    }
    if as_srt:
        # request raw SRT from the API
        data["srt"] = json.dumps({"format": "srt"})

    with open(file_path, "rb") as fp:
        resp = requests.post(
            ELEVENLABS_STT_URL,
            headers=headers,
            files={"file": fp},
            data=data,
        )
    resp.raise_for_status()

    return resp.text if as_srt else resp.json()

def get_thread_ts(file_info, channel_id):
    # Extract the Slack thread_ts from the file info.
    shares = file_info.get("shares", {})
    thread_ts = None
    if "public" in shares and channel_id in shares["public"]:
        thread_ts = shares["public"][channel_id][0].get("ts")
    elif "private" in shares and channel_id in shares["private"]:
        thread_ts = shares["private"][channel_id][0].get("ts")
    return thread_ts

def get_google_drive_service():
    """Authenticate and return a Google Drive service object."""
    creds = None
    try:
        if GOOGLE_CREDENTIALS_JSON:
            logging.info("Attempting to create Google Drive service from GOOGLE_CREDENTIALS_JSON env var.")
            creds_json = json.loads(GOOGLE_CREDENTIALS_JSON)
            creds = service_account.Credentials.from_service_account_info(
                creds_json,
                scopes=['https://www.googleapis.com/auth/drive']
            )
        elif GOOGLE_APPLICATION_CREDENTIALS:
            logging.info(f"Attempting to create Google Drive service with credentials file: {GOOGLE_APPLICATION_CREDENTIALS}")
            creds = service_account.Credentials.from_service_account_file(
                GOOGLE_APPLICATION_CREDENTIALS,
                scopes=['https://www.googleapis.com/auth/drive']
            )
        
        if not creds:
            logging.error("Google Drive credentials are not configured. Please set either GOOGLE_CREDENTIALS_JSON or GOOGLE_APPLICATION_CREDENTIALS.")
            return None

        service = build('drive', 'v3', credentials=creds)
        logging.info("Successfully created Google Drive service")
        return service
    except json.JSONDecodeError:
        logging.error("Failed to parse GOOGLE_CREDENTIALS_JSON. Make sure it's valid JSON.")
        return None
    except Exception as e:
        logging.error(f"Failed to create Google Drive service: {e}")
        return None

def get_or_create_shared_drive(service, drive_name="Transponster"):
    """Get the existing Transponster shared drive."""
    try:
        # List all shared drives
        drives = service.drives().list().execute()
        
        # Look for existing Transponster drive
        for drive in drives.get('drives', []):
            if drive['name'] == drive_name:
                logging.info(f"Found existing shared drive: {drive['name']} (ID: {drive['id']})")
                return drive['id']
        
        # If no Transponster drive found, return None
        logging.warning(f"No shared drive named '{drive_name}' found")
        return None
        
    except HttpError as error:
        logging.error(f"An error occurred with shared drive: {error}")
        return None

def find_or_create_folder(service, folder_name, parent_id=None):
    """Find a folder by name, or create it if it doesn't exist. Returns (id, webViewLink, created_boolean)."""
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    
    try:
        # First, try to find the folder in shared drives
        response = service.files().list(
            q=query, 
            spaces='drive', 
            fields='files(id, name, webViewLink)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        folders = response.get('files', [])
        
        if folders:
            return folders[0]['id'], folders[0]['webViewLink'], False # Existed
        
        # Create folder if not found
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            file_metadata['parents'] = [parent_id]
        
        folder = service.files().create(
            body=file_metadata, 
            fields='id, webViewLink',
            supportsAllDrives=True
        ).execute()
        return folder.get('id'), folder.get('webViewLink'), True # Created
    except HttpError as error:
        logging.error(f"An error occurred: {error}")
        return None, None, False

def upload_as_google_doc(service, file_name, file_content, folder_id):
    """Uploads text content as a .docx file to a specific folder."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        # Replace quotes before creating the document
        file_content = replace_quotes_with_ukrainian(file_content)

        # Create a Word document
        doc = Document()
        p = doc.add_paragraph(file_content)
        for run in p.runs:
            run.font.name = "Montserrat"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Montserrat')
        
        # Create a temporary file to upload
        temp_file_path = f"/tmp/{file_name}.docx"
        doc.save(temp_file_path)

        file_metadata = {
            'name': f"{file_name}.docx",
            'parents': [folder_id],
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        media = MediaFileUpload(temp_file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        
        file = service.files().create(
            body=file_metadata, 
            media_body=media, 
            fields='id, webViewLink',
            supportsAllDrives=True
        ).execute()
        
        # Clean up the temporary file
        os.remove(temp_file_path)
        
        return file.get('webViewLink')
    except HttpError as error:
        logging.error(f"An error occurred during file upload: {error}")
        return None
    except ImportError:
        logging.error("python-docx library not available, falling back to .txt file")
        # Fallback to .txt if python-docx is not available
        temp_file_path = f"/tmp/{file_name}.txt"
        with open(temp_file_path, "w", encoding="utf-8") as f:
            f.write(file_content)

        file_metadata = {
            'name': f"{file_name}.txt",
            'parents': [folder_id],
            'mimeType': 'text/plain'
        }
        
        media = MediaFileUpload(temp_file_path, mimetype='text/plain', resumable=True)

        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink',
            supportsAllDrives=True
        ).execute()

        # Clean up the temporary file
        os.remove(temp_file_path)

        return file.get('webViewLink')


def parse_srt_content(srt_text: str) -> list[dict]:
    """
    Parse SRT content and extract subtitle entries.
    Returns a list of dicts with 'index', 'timestamp', and 'text' keys.
    """
    entries = []
    blocks = re.split(r'\n\n+', srt_text.strip())

    for block in blocks:
        lines = block.strip().split('\n')
        if len(lines) >= 3:
            try:
                index = int(lines[0].strip())
                timestamp = lines[1].strip()
                text = '\n'.join(lines[2:]).strip()
                entries.append({
                    'index': index,
                    'timestamp': timestamp,
                    'text': text
                })
            except ValueError:
                continue

    return entries


TRANSLATION_BATCH_SIZE = 20
TRANSLATION_MAX_CONCURRENT = 8


def _make_translation_schema(n: int) -> dict:
    """Generate a JSON schema that enforces exactly n translated items."""
    return {
        "type": "json_schema",
        "json_schema": {
            "name": "translation_response",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {
                    "translated": {
                        "type": "array",
                        "items": {"type": "string"},
                        "minItems": n,
                        "maxItems": n
                    }
                },
                "required": ["translated"],
                "additionalProperties": False
            }
        }
    }


TRANSLATION_SYSTEM_PROMPT = """You are a subtitle translator. You will receive a JSON array of subtitle lines.

CRITICAL RULES:
1. Translate each line to English
2. Return EXACTLY the same number of items as you receive
3. NEVER merge lines together - each input line must produce exactly one output line
4. NEVER split lines - one input = one output
5. Preserve the order exactly

If a line is just a word like "Yes" or "No", translate it as a single word. Do not combine with adjacent lines."""


CLEANUP_SYSTEM_PROMPT = """Ти редактор усного мовлення. Твоя єдина задача — прибрати словесне сміття з розшифровки, НЕ змінюючи зміст.

ПРИБЕРИ:
- Звуки вагання: "еее", "е", "ммм", "хм", "и-и", "а-а"
- Слова-паразити: "ну", "от", "вот", "типу", "коротше", "скажем так", "як би", "там" (коли не вказує на місце)
- Зайві повтори слів підряд
- Русизми замінюй на українські відповідники: "получається"→"виходить", "вобще"→"взагалі", "надо"→"потрібно", "єтот/етот"→"цей", "конєчно"→"звісно"

НЕ ЧІПАЙ:
- Зміст і факти — жодних додавань чи вигадок
- Імена, назви, терміни, цифри
- Авторський стиль мовця (якщо людина каже "блін" — залиш)
- Структуру речень (не переписуй, лише чисти)

Якщо сегмент складається лише з "Угу", "Так", "Ага" — залиш як є.

Поверни JSON з полем "cleaned" — масив очищених рядків у тому ж порядку, що й вхідні."""


async def _translate_single_line(client, text: str) -> str:
    """Translate a single line - used as fallback."""
    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a translator. Translate the user's text to English. Output ONLY the translation, nothing else."},
            {"role": "user", "content": text}
        ],
        temperature=0
    )
    return response.choices[0].message.content.strip()


TRANSLATION_BATCH_TIMEOUT = 30  # seconds


async def _translate_batch(client, batch: list[str], batch_num: int, total_batches: int) -> list[str]:
    """
    Translate a batch of lines using structured outputs with length constraints.
    Returns translated lines, falling back to line-by-line if validation fails.
    """
    n = len(batch)
    schema = _make_translation_schema(n)

    for attempt in range(2):  # Try twice before fallback
        try:
            response = await asyncio.wait_for(
                client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": TRANSLATION_SYSTEM_PROMPT},
                        {"role": "user", "content": json.dumps(batch)}
                    ],
                    response_format=schema,
                    temperature=0
                ),
                timeout=TRANSLATION_BATCH_TIMEOUT
            )

            result = json.loads(response.choices[0].message.content)
            translated = result.get("translated", [])

            # Validate length matches (should always pass with minItems/maxItems, but safety check)
            if len(translated) == n:
                logging.info(f"[translation] Batch {batch_num}/{total_batches} OK ({n} lines)")
                return translated
            else:
                logging.warning(f"[translation] Batch {batch_num}/{total_batches} length mismatch: expected {n}, got {len(translated)} (attempt {attempt + 1})")

        except asyncio.TimeoutError:
            logging.warning(f"[translation] Batch {batch_num}/{total_batches} timeout after {TRANSLATION_BATCH_TIMEOUT}s (attempt {attempt + 1})")

        except Exception as e:
            logging.warning(f"[translation] Batch {batch_num}/{total_batches} error (attempt {attempt + 1}): {e}")

    # Fallback to line-by-line for this batch
    logging.info(f"[translation] Batch {batch_num}/{total_batches} falling back to line-by-line ({n} lines)")
    fallback_results = []
    for text in batch:
        translated = await _translate_single_line(client, text)
        fallback_results.append(translated)
    return fallback_results


async def translate_texts_with_openai(texts: list[str]) -> list[str]:
    """
    Translate a list of texts to English using OpenAI API.
    Uses batched requests with structured outputs for efficiency.
    Falls back to line-by-line translation if a batch fails validation.
    Returns a list of translated texts in the same order.
    """
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API_KEY is not configured")

    if not texts:
        return []

    client = AsyncOpenAI(api_key=OPENAI_API_KEY)

    # Split into batches
    batches = [texts[i:i + TRANSLATION_BATCH_SIZE] for i in range(0, len(texts), TRANSLATION_BATCH_SIZE)]
    total_batches = len(batches)

    logging.info(f"[translation] Starting translation of {len(texts)} lines in {total_batches} batches")

    # Process batches with controlled concurrency
    semaphore = asyncio.Semaphore(TRANSLATION_MAX_CONCURRENT)

    async def process_with_semaphore(batch, batch_num):
        async with semaphore:
            return await _translate_batch(client, batch, batch_num, total_batches)

    # Run all batches concurrently (limited by semaphore)
    tasks = [process_with_semaphore(batch, i + 1) for i, batch in enumerate(batches)]
    batch_results = await asyncio.gather(*tasks)

    # Flatten results
    translations = []
    for result in batch_results:
        translations.extend(result)

    logging.info(f"[translation] Completed translating {len(texts)} lines")
    return translations


# --- Cleanup (filler word removal) ---

def _make_cleanup_schema(n: int) -> dict:
    """Generate a JSON schema that enforces exactly n cleaned items."""
    return {
        "type": "json_schema",
        "json_schema": {
            "name": "cleanup_response",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {
                    "cleaned": {
                        "type": "array",
                        "items": {"type": "string"},
                        "minItems": n,
                        "maxItems": n
                    }
                },
                "required": ["cleaned"],
                "additionalProperties": False
            }
        }
    }


async def _cleanup_single_line(client, text: str) -> str:
    """Clean up a single line - used as fallback."""
    response = await client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Прибери слова-паразити (ну, еее, типу, коротше, скажем так) та русизми з тексту. Не змінюй зміст. Поверни лише очищений текст."},
            {"role": "user", "content": text}
        ],
        temperature=0
    )
    return response.choices[0].message.content.strip()


async def _cleanup_batch(client, batch: list[str], batch_num: int, total_batches: int) -> list[str]:
    """
    Clean up a batch of lines using structured outputs.
    Returns cleaned lines, falling back to line-by-line if validation fails.
    """
    n = len(batch)
    schema = _make_cleanup_schema(n)

    for attempt in range(2):
        try:
            response = await asyncio.wait_for(
                client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": CLEANUP_SYSTEM_PROMPT},
                        {"role": "user", "content": json.dumps(batch, ensure_ascii=False)}
                    ],
                    response_format=schema,
                    temperature=0
                ),
                timeout=TRANSLATION_BATCH_TIMEOUT
            )

            result = json.loads(response.choices[0].message.content)
            cleaned = result.get("cleaned", [])

            if len(cleaned) == n:
                logging.info(f"[cleanup] Batch {batch_num}/{total_batches} OK ({n} lines)")
                return cleaned
            else:
                logging.warning(f"[cleanup] Batch {batch_num}/{total_batches} length mismatch: expected {n}, got {len(cleaned)} (attempt {attempt + 1})")

        except asyncio.TimeoutError:
            logging.warning(f"[cleanup] Batch {batch_num}/{total_batches} timeout (attempt {attempt + 1})")

        except Exception as e:
            logging.warning(f"[cleanup] Batch {batch_num}/{total_batches} error (attempt {attempt + 1}): {e}")

    # Fallback to line-by-line
    logging.info(f"[cleanup] Batch {batch_num}/{total_batches} falling back to line-by-line ({n} lines)")
    fallback_results = []
    for text in batch:
        cleaned = await _cleanup_single_line(client, text)
        fallback_results.append(cleaned)
    return fallback_results


async def clean_texts_with_openai(texts: list[str]) -> list[str]:
    """
    Clean up a list of Ukrainian texts by removing filler words.
    Uses batched requests with structured outputs for efficiency.
    Returns a list of cleaned texts in the same order.
    """
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API_KEY is not configured")

    if not texts:
        return []

    client = AsyncOpenAI(api_key=OPENAI_API_KEY)

    batches = [texts[i:i + TRANSLATION_BATCH_SIZE] for i in range(0, len(texts), TRANSLATION_BATCH_SIZE)]
    total_batches = len(batches)

    logging.info(f"[cleanup] Starting cleanup of {len(texts)} lines in {total_batches} batches")

    semaphore = asyncio.Semaphore(TRANSLATION_MAX_CONCURRENT)

    async def process_with_semaphore(batch, batch_num):
        async with semaphore:
            return await _cleanup_batch(client, batch, batch_num, total_batches)

    tasks = [process_with_semaphore(batch, i + 1) for i, batch in enumerate(batches)]
    batch_results = await asyncio.gather(*tasks)

    cleaned_texts = []
    for result in batch_results:
        cleaned_texts.extend(result)

    logging.info(f"[cleanup] Completed cleaning {len(texts)} lines")
    return cleaned_texts


def rebuild_srt_with_translations(entries: list[dict], translations: list[str]) -> str:
    """
    Rebuild SRT content with translated texts.
    Takes parsed entries and list of translated texts.
    """
    output_lines = []

    for i, entry in enumerate(entries):
        output_lines.append(str(entry['index']))
        output_lines.append(entry['timestamp'])
        output_lines.append(translations[i] if i < len(translations) else entry['text'])
        output_lines.append('')  # Blank line between entries

    return '\n'.join(output_lines)


TRANSCRIPT_HEADER_PATTERN = re.compile(
    r'\d{2}:\d{2}:\d{2},\d{3} --> \d{2}:\d{2}:\d{2},\d{3} - \[.+\]'
)


def parse_transcript_content(txt_text: str) -> list[dict]:
    """
    Parse a Transponster .txt transcript into entries.
    Each entry has 'header' (timestamp + speaker line) and 'text' (spoken content).
    Returns empty list if no valid entries found.
    """
    lines = txt_text.split('\n')
    entries = []
    current_header = None
    current_text_lines = []

    for line in lines:
        if TRANSCRIPT_HEADER_PATTERN.match(line.strip()):
            # Save previous entry if exists
            if current_header is not None:
                text = '\n'.join(current_text_lines).strip()
                if text:
                    entries.append({'header': current_header, 'text': text})
            current_header = line.strip()
            current_text_lines = []
        elif current_header is not None:
            current_text_lines.append(line)

    # Save last entry
    if current_header is not None:
        text = '\n'.join(current_text_lines).strip()
        if text:
            entries.append({'header': current_header, 'text': text})

    return entries


def rebuild_transcript_with_translations(entries: list[dict], translations: list[str]) -> str:
    """
    Rebuild a .txt transcript using original headers and translated text.
    """
    output_lines = []
    for i, entry in enumerate(entries):
        output_lines.append(entry['header'])
        output_lines.append('')
        output_lines.append(translations[i] if i < len(translations) else entry['text'])
        output_lines.append('')
    return '\n'.join(output_lines)


def update_docx_with_content(service, file_id: str, content: str, heading_text: str) -> str | None:
    """
    Download existing .docx from Drive, append content as a new page with heading,
    re-upload in place. Returns webViewLink or None.
    """
    import io
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    try:
        # Download existing docx
        request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
        buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(buffer, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()

        buffer.seek(0)
        doc = Document(buffer)

        # Add page break and heading
        doc.add_page_break()
        heading = doc.add_heading(heading_text, level=1)
        for run in heading.runs:
            run.font.name = "Montserrat"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Montserrat')

        # Add content
        p = doc.add_paragraph(content)
        for run in p.runs:
            run.font.name = "Montserrat"
            run.font.size = Pt(14)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Montserrat')

        # Save to buffer and re-upload
        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        out_buffer.seek(0)

        media = MediaIoBaseUpload(
            out_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )
        updated = service.files().update(
            fileId=file_id,
            media_body=media,
            fields='webViewLink',
            supportsAllDrives=True
        ).execute()

        return updated.get('webViewLink')

    except Exception as e:
        logging.error(f"Failed to update docx {file_id}: {e}", exc_info=True)
        return None


def update_docx_with_translation(service, file_id: str, translated_content: str) -> str | None:
    """Append English translation to a Drive docx."""
    return update_docx_with_content(service, file_id, translated_content, "English version")


def update_docx_with_cleanup(service, file_id: str, cleaned_content: str) -> str | None:
    """Append cleaned version to a Drive docx."""
    return update_docx_with_content(service, file_id, cleaned_content, "Вичищена версія")